import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class DocxIncrementalGenerator:
    """
    DOCX 增量生成器
    基于现有文档模板，增量添加内容
    """

    def __init__(self, template_path=None):
        """
        初始化生成器

        Args:
            template_path: 模板文档路径，如果为None则创建新文档
        """
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()

        # 设置中文字体（如果需要）
        self.set_chinese_font()

    def set_chinese_font(self, font_name='微软雅黑'):
        """
        设置文档的中文字体
        """
        try:
            self.doc.styles['Normal'].font.name = font_name
            self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        except Exception as e:
            print(f"字体设置失败: {e}")

    def add_heading(self, text, level=1):
        """
        添加标题

        Args:
            text: 标题文本
            level: 标题级别 (1-9)
        """
        self.doc.add_heading(text, level=level)
        return self

    def add_paragraph(self, text, style=None, bold=False, italic=False, font_size=12):
        """
        添加段落

        Args:
            text: 段落文本
            style: 段落样式
            bold: 是否加粗
            italic: 是否斜体
            font_size: 字体大小
        """
        paragraph = self.doc.add_paragraph(text, style=style)

        # 设置字体格式
        if bold or italic or font_size != 12:
            for run in paragraph.runs:
                run.bold = bold
                run.italic = italic
                if font_size != 12:
                    run.font.size = Pt(font_size)

        return self

    def add_table(self, data, headers=None, style='Light Grid Accent 1'):
        """
        添加表格

        Args:
            data: 表格数据（二维列表）
            headers: 表头列表
            style: 表格样式
        """
        if headers:
            # 包含表头
            table_data = [headers] + data
        else:
            table_data = data

        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style

        # 填充数据
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                table.cell(i, j).text = str(cell_data)

        return self

    def add_list(self, items, ordered=False):
        """
        添加列表

        Args:
            items: 列表项
            ordered: 是否有序列表
        """
        list_style = 'List Number' if ordered else 'List Bullet'

        for item in items:
            self.doc.add_paragraph(item, style=list_style)

        return self

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
        return self

    def add_image(self, image_path, width=6):
        """
        添加图片

        Args:
            image_path: 图片路径
            width: 图片宽度（英寸）
        """
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=Inches(width))
        else:
            print(f"图片文件不存在: {image_path}")

        return self

    def save(self, output_path):
        """
        保存文档

        Args:
            output_path: 输出路径
        """
        try:
            self.doc.save(output_path)
            print(f"文档已保存: {output_path}")
            return True
        except Exception as e:
            print(f"保存失败: {e}")
            return False


# 使用示例
def example_usage():
    """使用示例"""

    # 创建生成器实例（可以传入模板路径）
    generator = DocxIncrementalGenerator()

    # 添加内容
    generator.add_heading('Python DOCX 增量生成示例', level=1)

    generator.add_paragraph('这是一个使用 python-docx 库进行增量生成的示例文档。')

    # 添加表格
    table_data = [
        ['Python', '3.9', '高级编程语言'],
        ['Java', '17', '面向对象语言'],
        ['JavaScript', 'ES2022', '脚本语言']
    ]
    headers = ['语言', '版本', '描述']
    generator.add_table(table_data, headers=headers)

    # 添加列表
    items = ['文档生成', '格式保持', '增量添加', '模板支持']
    generator.add_list(items, ordered=False)

    # 添加带格式的段落
    generator.add_paragraph('这是加粗的文本。', bold=True, font_size=14)
    generator.add_paragraph('这是斜体的文本。', italic=True)

    # 保存文档
    output_file = f"generated_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    generator.save(output_file)


# 高级功能：基于模板的增量生成
def template_based_generation():
    """基于模板的增量生成示例"""

    # 假设有一个模板文件
    template_path = "template.docx"  # 替换为实际模板路径

    if os.path.exists(template_path):
        generator = DocxIncrementalGenerator(template_path)
        generator.add_heading('新增内容', level=2)
        generator.add_paragraph('这是在模板基础上新增的内容。')

        output_file = f"template_based_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        generator.save(output_file)
    else:
        print("模板文件不存在，创建新文档示例")
        example_usage()


if __name__ == "__main__":
    # 运行示例
    example_usage()

    # 如果需要基于模板生成，取消下面的注释
    # template_based_generation()